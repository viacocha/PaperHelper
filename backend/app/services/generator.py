from __future__ import annotations

from pathlib import Path

from docx import Document

from app.models.schemas import GeneratedPaper, StandardMatch
from app.services.standards import Standard, StandardLibrary

MIN_GENERATED_WORDS = 2500


def generate_paper(
    project_background: str,
    standard_id: str,
    library: StandardLibrary,
    output_dir: Path,
) -> tuple[GeneratedPaper, Path]:
    standard = _find_standard(standard_id, library)
    title = f"论信息系统项目的{standard.name}"
    if standard.category == "performance_domain":
        title = f"论信息系统项目{standard.name}"

    paragraphs = _build_paragraphs(project_background.strip(), standard)
    paragraphs = _ensure_minimum_length(paragraphs, standard, MIN_GENERATED_WORDS)
    content = "\n\n".join(paragraphs)
    word_count = _count_words(content)
    report_name = f"{standard.name}+论文初稿.docx"
    output_path = output_dir / report_name
    _write_docx(title, paragraphs, output_path)

    return (
        GeneratedPaper(
            title=title,
            standard=StandardMatch(
                standard_id=standard.id,
                standard_name=standard.name,
                category=standard.category,  # type: ignore[arg-type]
                confidence=1.0,
            ),
            content=content,
            paragraphs=paragraphs,
            word_count=word_count,
            minimum_word_count=MIN_GENERATED_WORDS,
            generated_report_name=report_name,
        ),
        output_path,
    )


def _find_standard(standard_id: str, library: StandardLibrary) -> Standard:
    for standard in library.all():
        if standard.id == standard_id:
            return standard
    raise ValueError("请选择有效的论文题型。")


def _build_paragraphs(project_background: str, standard: Standard) -> list[str]:
    process_text = "、".join(standard.required_processes[:6])
    artifact_text = "、".join(standard.required_artifacts) or "相关管理计划和项目文件"
    question_text = "、".join(standard.question_points)
    category_text = "绩效域" if standard.category == "performance_domain" else "知识域"
    process_paragraphs = _process_paragraphs(standard)

    paragraphs = [
        f"【题目】{_title_for(standard)}",
        (
            f"【摘要】本文以我负责的一个信息系统建设项目为背景，围绕{standard.name}这一{category_text}论文主题，"
            f"从项目背景、管理思路、重点过程、实际问题、解决措施和实施效果等方面展开论述。"
            f"项目实施过程中，我坚持以目标为导向，以过程控制为抓手，结合{process_text}等管理活动，"
            f"通过{artifact_text}等成果沉淀管理过程，最终保障项目顺利推进并通过验收。"
        ),
        (
            f"{project_background}。在该项目中，我担任乙方项目经理，全面负责项目启动、规划、执行、监控和收尾工作。"
            f"项目具有业务部门多、接口协同复杂、需求调整频繁、上线窗口固定等特点。结合这些特点，我认为{standard.name}是保障项目目标实现的重要管理内容。"
            f"本文将结合该项目实践，围绕{question_text}展开论述。"
        ),
        (
            f"{standard.name}不是单纯的理论概念，而是需要在真实项目场景中持续落地的管理活动。"
            f"在本项目中，我将其与项目目标、交付成果、干系人诉求和项目约束条件结合起来，重点从{process_text}等方面开展管理，"
            f"并通过{artifact_text}等工具和文档进行跟踪、控制和复盘。"
            f"我的总体思路是先建立统一的管理规则，再把规则分解到计划、会议、评审、台账和验收活动中，最后通过持续监控和闭环处理保证管理措施真正发挥作用。"
        ),
    ]

    paragraphs.extend(process_paragraphs)
    paragraphs.append(
        f"为提升{standard.name}的实践深度，我重点使用了{artifact_text}等工具。"
        f"例如，我会将关键事项分解为可跟踪条目，明确来源、责任人、处理状态、完成时间和验收标准，并在项目周例会和阶段评审会上持续更新。"
        f"对于影响范围较大的事项，我要求团队形成书面记录，说明问题背景、分析依据、备选方案、最终决策和后续责任人。"
        f"这样既能让团队及时发现偏差，也便于在项目监控过程中形成闭环，避免只开会不落地、只记录不跟踪的情况。"
    )

    paragraphs.extend([
        (
            f"在项目执行过程中，{standard.name}也遇到了一些管理难点。"
            f"例如，部分干系人对交付边界、实施节奏或验收标准的理解不一致，导致团队一度出现返工和沟通成本上升；"
            f"个别外部接口单位响应不及时，也对后续联调和测试安排造成影响。"
            f"为此，我组织相关干系人召开专题协调会，重新确认关键事项，并通过问题清单和变更记录持续跟踪处理结果。"
            f"对于确需调整的事项，我要求先分析对范围、进度、成本、质量和风险的影响，再提交变更评审，避免口头承诺直接进入开发。"
            f"经过以上措施，项目管理过程逐步回到可控状态。"
        ),
        (
            f"在监控阶段，我坚持以数据和事实为依据，对{standard.name}相关事项进行持续检查。"
            f"对已经完成的工作，我组织团队进行阶段性复盘；对仍存在偏差的事项，我及时协调资源、调整计划并向关键干系人同步。"
            f"我还将关键指标纳入项目状态报告，包括计划完成率、问题关闭率、缺陷修复率、变更处理周期和验收通过情况等。"
            f"通过这种方式，项目团队能够及时识别问题、快速采取措施，并保证最终交付成果符合客户要求。"
        ),
        (
            f"项目进入收尾阶段后，我组织团队对{standard.name}相关过程进行总结。"
            f"一方面，我与甲方代表、业务骨干、测试人员和运维人员一起确认交付物清单、培训记录、试运行问题和验收材料，确保项目成果能够被用户真正接收；"
            f"另一方面，我要求项目组复盘管理过程中的有效做法和不足之处，将典型问题、处理措施和经验教训沉淀到组织过程资产中。"
            f"这些收尾工作不仅有助于项目顺利验收，也为后续类似项目提供了可复用的管理经验。"
        ),
        (
            f"最终，在项目团队和各方干系人的共同努力下，项目按计划完成主要建设任务并顺利通过验收。"
            f"通过本项目实践，我进一步认识到，{standard.name}的关键不在于机械背诵流程，而在于结合项目场景，将理论方法转化为具体管理动作。"
            f"今后在类似信息系统项目中，我将继续坚持理论联系实际，重视前期规划、过程控制、问题闭环和经验复盘，不断提升项目管理水平。"
        ),
    ])
    return paragraphs


def _process_paragraphs(standard: Standard) -> list[str]:
    transitions = ["首先", "其次", "再次", "然后", "同时", "最后"]
    paragraphs: list[str] = []
    for index, process in enumerate(standard.required_processes[:6]):
        transition = transitions[index] if index < len(transitions) else "此外"
        paragraphs.append(
            f"{transition}，在{process}过程中，我结合项目背景和客户诉求，组织项目团队识别本过程的关键输入、约束条件和预期输出。"
            f"我没有把{process}简单理解为编写一份文档，而是将其转化为可执行的管理动作：明确参与人员、确认工作方法、定义检查标准，并把输出成果纳入项目整体计划。"
            f"针对项目中可能出现的沟通不充分、责任边界不清或执行偏差等问题，我通过专题会议、评审机制和项目文件跟踪的方式进行管理。"
            f"在具体执行中，我要求团队将管理动作落实到责任人、时间节点和可检查成果上，确保{process}能够支撑项目后续交付。"
        )
    return paragraphs


def _ensure_minimum_length(paragraphs: list[str], standard: Standard, minimum_word_count: int) -> list[str]:
    if _count_words("\n\n".join(paragraphs)) >= minimum_word_count:
        return paragraphs

    artifact_text = "、".join(standard.required_artifacts) or "项目管理计划、问题清单和状态报告"
    extra_templates = [
        (
            f"为了避免论文主题与项目实践脱节，我在项目管理过程中始终围绕{standard.name}建立闭环。"
            f"在每一次重要评审前，我都会检查相关输入是否完整、责任人是否明确、输出成果是否可以被验证；评审后，我再把遗留问题纳入台账进行跟踪。"
            f"如果某项工作没有形成明确结论，我不会简单认为会议已经完成，而是要求责任人补充分析材料，并在下一次例会上确认处理结果。"
            f"这种做法使{standard.name}从一次性活动变成持续性的管理机制。"
        ),
        (
            f"在工具和文档方面，我重点维护{artifact_text}。"
            f"这些文档不是为了应付检查，而是为了支撑项目决策和团队协同。"
            f"例如，当业务部门提出新增需求或调整验收标准时，我会先检查相关记录，判断该事项是否已经纳入计划、是否影响既有承诺、是否需要提交变更。"
            f"通过文档化管理，项目团队能够减少口头沟通带来的理解偏差，也便于在后续验收和审计时说明管理依据。"
        ),
        (
            f"在团队管理上，我也将{standard.name}与沟通、质量、进度和风险等管理内容结合起来。"
            f"对于跨部门协作事项，我会明确牵头人和配合人，要求重要信息在统一渠道发布，避免同一问题被不同人员重复解释。"
            f"对于影响关键路径或重要交付物的问题，我会提高跟踪频率，必要时协调甲方项目负责人共同推动。"
            f"这种联动式管理让项目成员能够理解自己的工作对整体目标的影响，从而提高执行效率。"
        ),
        (
            f"从结果来看，项目虽然在实施过程中出现过需求澄清、接口联调和验收材料准备等方面的压力，但由于{standard.name}相关措施落实较早，问题暴露后能够及时被记录、分析和处理。"
            f"项目组没有等到最终验收时才集中解决问题，而是在阶段性交付和试运行过程中持续消化偏差。"
            f"这也说明，在信息系统项目中，项目经理必须把管理标准转化为日常行动，才能真正提高项目的可控性和成功率。"
        ),
    ]

    extended = paragraphs[:]
    while _count_words("\n\n".join(extended)) < minimum_word_count:
        extended.extend(extra_templates)
    return extended


def _count_words(text: str) -> int:
    compact = "".join(char for char in text if not char.isspace())
    return len(compact)


def _title_for(standard: Standard) -> str:
    if standard.category == "performance_domain":
        return f"论信息系统项目{standard.name}"
    return f"论信息系统项目的{standard.name}"


def _write_docx(title: str, paragraphs: list[str], output_path: Path) -> None:
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs[1:]:
        document.add_paragraph(paragraph)
    document.save(output_path)
