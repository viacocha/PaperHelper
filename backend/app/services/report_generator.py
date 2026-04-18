from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from app.models.schemas import CompareResult, ReviewResult

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"


def generate_report(review: ReviewResult, output_path: Path) -> Path:
    document = Document()
    _add_title(document, f"{review.filename} 批改报告")
    _add_heading(document, "一、批改总评")
    document.add_paragraph(f"总分：{review.total_score} / 75")
    document.add_paragraph(f"通过参考线：{review.pass_score}")
    document.add_paragraph(f"通过风险判断：{_label(review.pass_probability)}")
    document.add_paragraph(f"匹配题型：{review.standard.standard_name}（{review.standard.category}，置信度 {review.standard.confidence}）")
    document.add_paragraph(f"总体结论：{review.summary}")

    _add_heading(document, "二、分项评分")
    for dimension in review.dimensions:
        document.add_paragraph(
            f"{dimension.name}：{dimension.score} / {dimension.max_score}。{dimension.summary}",
            style="List Bullet",
        )

    _add_heading(document, "三、题型评分卡")
    for item in review.topic_scorecard:
        document.add_paragraph(
            f"[{item.status.upper()}] {item.title}：{item.score} / {item.max_score}。{item.summary}",
            style="List Bullet",
        )

    _add_heading(document, "四、主要问题与修改建议")
    if review.issues:
        for item in review.issues:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"[{item.action_priority.upper()}/{item.severity.upper()}] {item.title}：").bold = True
            paragraph.add_run(item.details + " ")
            paragraph.add_run("建议：").bold = True
            paragraph.add_run(item.suggestion)
    else:
        document.add_paragraph("未检测到明显高风险问题，但仍建议结合题目子问逐项复核。")

    _add_heading(document, "五、修改优先级队列")
    _add_issue_group(document, "必须补写", review.must_fix)
    _add_issue_group(document, "建议补强", review.should_fix)
    _add_issue_group(document, "可优化项", review.could_improve)

    _add_heading(document, "六、题型模板建议")
    for item in review.revision_templates:
        document.add_paragraph(item.title).runs[0].bold = True
        document.add_paragraph("用途：" + item.purpose)
        document.add_paragraph("适用场景：" + item.when_to_use)
        document.add_paragraph("补写结构：" + "；".join(item.outline))
        document.add_paragraph("示例写法：" + item.sample)

    _add_heading(document, "七、逐段建议")
    for paragraph_review in review.paragraph_reviews:
        document.add_paragraph(f"第 {paragraph_review.index} 段：{paragraph_review.excerpt}", style="List Bullet")
        if paragraph_review.strengths:
            document.add_paragraph("优点：" + "；".join(paragraph_review.strengths))
        if paragraph_review.issues:
            document.add_paragraph("问题：" + "；".join(paragraph_review.issues))
        if paragraph_review.suggestions:
            document.add_paragraph("建议：" + "；".join(paragraph_review.suggestions))

    _add_heading(document, "八、修改顺序建议")
    priorities = [
        "先补齐题目要求的子问、专属产物和关键过程。",
        "再强化项目经理视角、问题应对和结果成效。",
        "最后压缩空泛理论，润色衔接与专业术语表达。",
    ]
    for item in priorities:
        document.add_paragraph(item, style="List Number")

    document.save(output_path)
    return output_path


def generate_annotated_report(source_path: Path, review: ReviewResult, output_path: Path) -> Path:
    document = Document(source_path)
    _remove_existing_comment_markers(document)
    non_empty_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    comment_texts: list[str] = []

    title_paragraph = non_empty_paragraphs[0] if non_empty_paragraphs else document.add_paragraph(review.filename)
    _add_word_comment_marker(title_paragraph, 0)
    comment_texts.append(_overall_comment_text(review))

    for paragraph_review in review.paragraph_reviews:
        paragraph_index = paragraph_review.index - 1
        if paragraph_index >= len(non_empty_paragraphs):
            continue
        comment_text = _paragraph_comment_text(paragraph_review)
        if not comment_text:
            continue
        _add_word_comment_marker(non_empty_paragraphs[paragraph_index], len(comment_texts))
        comment_texts.append(comment_text)

    with NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_path = Path(temp_file.name)
    document.save(temp_path)

    _patch_comments_part(temp_path, output_path, comment_texts)
    temp_path.unlink(missing_ok=True)
    return output_path


def generate_compare_report(comparison: CompareResult, output_path: Path) -> Path:
    document = Document()
    _add_title(document, f"{comparison.revised.filename} 二次批改对比报告")

    _add_heading(document, "一、对比总评")
    document.add_paragraph(f"修改前总分：{comparison.original.total_score} / 75")
    document.add_paragraph(f"修改后总分：{comparison.revised.total_score} / 75")
    document.add_paragraph(f"分数变化：{comparison.score_delta:+.1f}")
    document.add_paragraph("通过风险变化：" + ("已变化" if comparison.pass_probability_changed else "未变化"))
    document.add_paragraph("总体结论：" + comparison.summary)

    _add_heading(document, "二、已修复问题")
    _add_issue_group(document, "已修复", comparison.fixed_issues)

    _add_heading(document, "三、仍未解决问题")
    _add_issue_group(document, "仍未解决", comparison.remaining_issues)

    _add_heading(document, "四、新增问题")
    _add_issue_group(document, "新增", comparison.new_issues)

    _add_heading(document, "五、下一轮修改建议")
    if comparison.remaining_issues:
        document.add_paragraph("优先处理仍未解决的问题，尤其是必须补写项和题目专属产物缺失。")
    if comparison.new_issues:
        document.add_paragraph("检查新增问题是否由压缩、改写或结构调整导致，避免修复旧问题时引入新缺陷。")
    if not comparison.remaining_issues and not comparison.new_issues:
        document.add_paragraph("当前未发现旧问题残留或新增问题，可进入表达润色和考场默写训练阶段。")

    document.save(output_path)
    return output_path


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _add_heading(document: Document, text: str) -> None:
    document.add_paragraph().add_run(text).bold = True


def _add_issue_group(document: Document, title: str, items) -> None:
    document.add_paragraph(title).runs[0].bold = True
    if not items:
        document.add_paragraph("当前无对应问题。")
        return
    for item in items:
        document.add_paragraph(f"{item.title}：{item.suggestion}", style="List Bullet")


def _label(value: str) -> str:
    return {
        "high": "较高通过概率",
        "medium": "接近通过，仍需修改",
        "low": "当前通过风险高",
    }[value]


def _add_word_comment_marker(paragraph, comment_id: int) -> str:  # type: ignore[no-untyped-def]
    if not paragraph.runs:
        paragraph.add_run(" ")

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    paragraph._p.insert(0, start)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    paragraph._p.append(end)

    reference_run = OxmlElement("w:r")
    reference = OxmlElement("w:commentReference")
    reference.set(qn("w:id"), str(comment_id))
    reference_run.append(reference)
    paragraph._p.append(reference_run)
    return str(comment_id)


def _remove_existing_comment_markers(document: Document) -> None:
    comment_tags = {
        qn("w:commentRangeStart"),
        qn("w:commentRangeEnd"),
        qn("w:commentReference"),
    }
    for element in list(document.element.iter()):
        if element.tag not in comment_tags:
            continue
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _overall_comment_text(review: ReviewResult) -> str:
    pass_text = {
        "high": "这篇整体基础不错，有较大机会通过",
        "medium": "这篇已经接近通过，但还需要把下面几处补扎实",
        "low": "这篇现在通过风险比较高，建议先按批注做一次比较大的修改",
    }[review.pass_probability]
    must_fix = "；".join(item.title for item in review.must_fix[:4])
    should_fix = "；".join(item.title for item in review.should_fix[:4])
    text = (
        f"总评：{pass_text}。我按阅卷要求估分大约 {review.total_score}/75，及格参考线是 {review.pass_score}。"
        f"这篇对应的题型是{review.standard.standard_name}。{review.summary}"
    )
    if must_fix:
        text += f"先改这些硬伤：{must_fix}。"
    if should_fix:
        text += f"如果时间够，再继续加强：{should_fix}。"
    text += "修改时不要只背概念，重点把题目要求的过程、工具或文件写进自己的项目场景里，写清楚怎么用、发现了什么、最后有什么效果。"
    return text


def _paragraph_comment_text(paragraph_review) -> str:  # type: ignore[no-untyped-def]
    parts: list[str] = []
    if paragraph_review.issues:
        parts.append("修改建议：" + "；".join(_humanize_comment(item) for item in paragraph_review.issues))
    if paragraph_review.suggestions:
        parts.append("建议你这样处理：" + "；".join(_humanize_comment(item) for item in paragraph_review.suggestions))
    if not parts:
        return ""
    if paragraph_review.strengths:
        parts.append("这一段保留的点：" + "；".join(_humanize_comment(item) for item in paragraph_review.strengths))
    return " ".join(parts)


def _humanize_comment(text: str) -> str:
    exact_replacements = {
        "本段缺少第一人称管理者视角。": "这一段还看不出你作为项目经理具体做了什么",
        "补充“我组织/我制定/我协调/我推动”等管理动作。": "建议补上“我组织、我制定、我协调、我推动”这类动作，让阅卷老师看到是你在管理项目",
        "本段缺少问题-措施-结果闭环。": "这里还缺一条完整的处理线：遇到了什么问题、你怎么处理、最后效果如何",
        "增加遇到的问题、采取的措施及最终效果。": "加上一个真实的小场景，把问题、措施和结果写完整",
        "本段内容偏短，信息密度不足。": "这一段太短了，阅卷时会显得内容比较空",
        "扩充本段中的项目事实、工具文档或结果。": "可以展开写项目事实、用到的工具文档，或者最后形成的结果",
        "本段提到了过程或工具，但缺少具体使用场景和数据佐证。": "这里虽然提到了过程或工具，但还没有写清楚具体怎么用，也缺少数据或效果来支撑",
        "补充工具如何制定、如何使用、发现了什么问题、如何纠正以及产生的效果。": "建议把工具的制定过程、使用过程、发现的问题、纠正办法和最后效果补出来",
        "包含项目事实信息。": "这一段有项目事实，可以保留",
        "体现了项目经理管理动作。": "这里能看出一些项目经理的管理动作",
        "命中了当前题型的关键术语。": "这里提到了这个题目的关键词",
    }
    if text in exact_replacements:
        return exact_replacements[text]

    replacements = {
        "本段": "这里",
        "缺少": "还没有写出",
        "补充": "补上",
        "增加": "加上",
        "扩充": "展开写一下",
        "命中了": "提到了",
        "当前题型": "这个题目",
        "关键术语": "关键词",
        "包含": "有",
        "体现了": "能看出",
        "信息密度不足": "内容还偏薄",
        "项目经理管理动作": "项目经理的管理动作",
        "问题-措施-结果闭环": "问题、处理措施和结果这一条线",
        "具体使用场景和数据佐证": "具体怎么用、用了以后有什么数据或效果",
    }
    result = text.strip().rstrip("。")
    for old, new in replacements.items():
        result = result.replace(old, new)
    return result


def _patch_comments_part(temp_path: Path, output_path: Path, comment_texts: list[str]) -> None:
    with ZipFile(temp_path, "r") as source:
        entries = {item.filename: source.read(item.filename) for item in source.infolist()}

    entries["word/comments.xml"] = _build_comments_xml(comment_texts)
    entries["word/_rels/document.xml.rels"] = _ensure_comments_relationship(
        entries.get("word/_rels/document.xml.rels", _empty_relationships_xml())
    )
    entries["[Content_Types].xml"] = _ensure_comments_content_type(entries["[Content_Types].xml"])

    with ZipFile(output_path, "w", ZIP_DEFLATED) as target:
        for name, data in entries.items():
            target.writestr(name, data)


def _build_comments_xml(comment_texts: list[str]) -> bytes:
    root = etree.Element(f"{{{WORD_NS}}}comments", nsmap={"w": WORD_NS})
    timestamp = datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    for index, text in enumerate(comment_texts):
        comment = etree.SubElement(root, f"{{{WORD_NS}}}comment")
        comment.set(f"{{{WORD_NS}}}id", str(index))
        comment.set(f"{{{WORD_NS}}}author", "王老师")
        comment.set(f"{{{WORD_NS}}}initials", "WL")
        comment.set(f"{{{WORD_NS}}}date", timestamp)
        paragraph = etree.SubElement(comment, f"{{{WORD_NS}}}p")
        run = etree.SubElement(paragraph, f"{{{WORD_NS}}}r")
        text_node = etree.SubElement(run, f"{{{WORD_NS}}}t")
        text_node.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _ensure_comments_relationship(raw_xml: bytes) -> bytes:
    root = etree.fromstring(raw_xml)
    for relationship in root:
        if relationship.get("Type") == COMMENTS_REL_TYPE:
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    existing_ids = {relationship.get("Id") for relationship in root}
    next_index = 1
    while f"rId{next_index}" in existing_ids:
        next_index += 1
    relationship = etree.SubElement(root, f"{{{RELS_NS}}}Relationship")
    relationship.set("Id", f"rId{next_index}")
    relationship.set("Type", COMMENTS_REL_TYPE)
    relationship.set("Target", "comments.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _ensure_comments_content_type(raw_xml: bytes) -> bytes:
    root = etree.fromstring(raw_xml)
    for override in root:
        if override.get("PartName") == "/word/comments.xml":
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    override = etree.SubElement(root, f"{{{CONTENT_TYPES_NS}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set("ContentType", COMMENTS_CONTENT_TYPE)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _empty_relationships_xml() -> bytes:
    root = etree.Element(f"{{{RELS_NS}}}Relationships", nsmap={None: RELS_NS})
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
